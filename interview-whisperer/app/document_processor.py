"""
Document Processor for Interview Whisperer

Extracts text from PDF, DOCX, TXT, and MD files, chunks them intelligently,
creates embeddings using Ollama, and stores them in ChromaDB.
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Optional, Callable, Tuple
from datetime import datetime
import hashlib

# Document parsing libraries
import PyPDF2
from docx import Document

# Vector database
import chromadb
from chromadb.config import Settings

# LLM integration
try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False
    print("Warning: ollama package not installed. Install with: pip install ollama")


class DocumentProcessor:
    """
    Process documents and store them in ChromaDB with embeddings.

    Supports: PDF, DOCX, TXT, MD files
    Uses: Ollama (nomic-embed-text) for embeddings
    Stores: ChromaDB for vector storage
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md'}
    EMBEDDING_MODEL = 'nomic-embed-text'
    COLLECTION_NAME = 'interview_context'

    def __init__(self, documents_dir: str, db_path: str):
        """
        Initialize the DocumentProcessor.

        Args:
            documents_dir: Path to directory containing documents
            db_path: Path to ChromaDB storage location
        """
        self.documents_dir = Path(documents_dir)
        self.db_path = Path(db_path)

        # Setup logging
        self._setup_logging()

        # Ensure directories exist
        self.documents_dir.mkdir(parents=True, exist_ok=True)
        self.db_path.mkdir(parents=True, exist_ok=True)

        # Initialize ChromaDB
        try:
            self.client = chromadb.PersistentClient(
                path=str(self.db_path),
                settings=Settings(anonymized_telemetry=False)
            )
            self.collection = self.client.get_or_create_collection(
                name=self.COLLECTION_NAME,
                metadata={"description": "Interview preparation context documents"}
            )
            self.logger.info(f"ChromaDB initialized at {self.db_path}")
        except Exception as e:
            self.logger.error(f"Failed to initialize ChromaDB: {e}")
            raise

    def _setup_logging(self):
        """Setup logging configuration."""
        log_dir = Path(self.db_path).parent / "logs"
        log_dir.mkdir(parents=True, exist_ok=True)

        log_file = log_dir / f"document_processor_{datetime.now().strftime('%Y%m%d')}.log"

        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)

    def scan_documents(self) -> List[Dict[str, str]]:
        """
        Scan documents directory for supported file types.

        Returns:
            List of dictionaries with file metadata:
            [{
                'path': str,
                'name': str,
                'extension': str,
                'size': int,
                'modified': datetime
            }]
        """
        documents = []

        try:
            for file_path in self.documents_dir.rglob('*'):
                if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                    documents.append({
                        'path': str(file_path),
                        'name': file_path.name,
                        'extension': file_path.suffix.lower(),
                        'size': file_path.stat().st_size,
                        'modified': datetime.fromtimestamp(file_path.stat().st_mtime)
                    })

            self.logger.info(f"Found {len(documents)} documents to process")
            return documents

        except Exception as e:
            self.logger.error(f"Error scanning documents: {e}")
            return []

    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from a file based on its type.

        Args:
            file_path: Path to the file

        Returns:
            Extracted text or None if extraction failed
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        try:
            if extension == '.pdf':
                return self._extract_from_pdf(path)
            elif extension == '.docx':
                return self._extract_from_docx(path)
            elif extension in {'.txt', '.md'}:
                return self._extract_from_text(path)
            else:
                self.logger.warning(f"Unsupported file type: {extension}")
                return None

        except Exception as e:
            self.logger.error(f"Failed to extract text from {file_path}: {e}")
            return None

    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = []

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                except Exception as e:
                    self.logger.warning(f"Failed to extract page {page_num} from {file_path.name}: {e}")

        return '\n'.join(text)

    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text.append(row_text)

        return '\n'.join(text)

    def _extract_from_text(self, file_path: Path) -> str:
        """Extract text from TXT or MD file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue

        self.logger.error(f"Failed to decode {file_path.name} with any encoding")
        return ""

    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """
        Split text into overlapping chunks, respecting sentence boundaries.

        Args:
            text: Text to chunk
            chunk_size: Target size in words
            overlap: Number of words to overlap between chunks

        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []

        # Clean text
        text = self._clean_text(text)

        # Split into sentences
        sentences = self._split_into_sentences(text)

        chunks = []
        current_chunk = []
        current_word_count = 0

        for sentence in sentences:
            sentence_words = len(sentence.split())

            # If adding this sentence exceeds chunk_size, save current chunk
            if current_word_count + sentence_words > chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))

                # Calculate overlap: keep last 'overlap' words
                overlap_words = ' '.join(current_chunk).split()[-overlap:]
                current_chunk = [' '.join(overlap_words)]
                current_word_count = len(overlap_words)

            current_chunk.append(sentence)
            current_word_count += sentence_words

        # Add remaining chunk
        if current_chunk:
            chunks.append(' '.join(current_chunk))

        self.logger.info(f"Created {len(chunks)} chunks from text")
        return chunks

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()

    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting (can be improved with NLTK if needed)
        sentence_endings = re.compile(r'([.!?])\s+')
        sentences = sentence_endings.split(text)

        # Recombine sentences with their punctuation
        result = []
        for i in range(0, len(sentences) - 1, 2):
            sentence = sentences[i] + (sentences[i + 1] if i + 1 < len(sentences) else '')
            if sentence.strip():
                result.append(sentence.strip())

        # Handle last sentence if it doesn't end with punctuation
        if len(sentences) % 2 == 1 and sentences[-1].strip():
            result.append(sentences[-1].strip())

        return result

    def create_embeddings(self, chunks: List[str]) -> Optional[List[List[float]]]:
        """
        Generate embeddings for text chunks using Ollama.

        Args:
            chunks: List of text chunks

        Returns:
            List of embeddings or None if generation failed
        """
        if not OLLAMA_AVAILABLE:
            self.logger.error("Ollama package not available")
            return None

        embeddings = []

        try:
            for i, chunk in enumerate(chunks):
                try:
                    response = ollama.embeddings(
                        model=self.EMBEDDING_MODEL,
                        prompt=chunk
                    )
                    embeddings.append(response['embedding'])

                except Exception as e:
                    self.logger.error(f"Failed to create embedding for chunk {i}: {e}")
                    # Return None to indicate failure
                    return None

            self.logger.info(f"Created {len(embeddings)} embeddings")
            return embeddings

        except Exception as e:
            self.logger.error(f"Ollama embedding generation failed: {e}")
            self.logger.error("Make sure Ollama is running and the model is installed:")
            self.logger.error(f"  ollama pull {self.EMBEDDING_MODEL}")
            return None

    def store_chunks(self, chunks: List[str], embeddings: List[List[float]],
                    metadata: Dict[str, str]) -> bool:
        """
        Store chunks and embeddings in ChromaDB.

        Args:
            chunks: List of text chunks
            embeddings: List of embeddings
            metadata: Metadata about the source file

        Returns:
            True if successful, False otherwise
        """
        if not chunks or not embeddings or len(chunks) != len(embeddings):
            self.logger.error("Chunks and embeddings length mismatch")
            return False

        try:
            # Generate unique IDs for each chunk
            file_hash = hashlib.md5(metadata['file_name'].encode()).hexdigest()[:8]
            ids = [f"{file_hash}_chunk_{i}" for i in range(len(chunks))]

            # Create metadata for each chunk
            metadatas = [
                {
                    'file_name': metadata['file_name'],
                    'file_type': metadata['file_type'],
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'processed_at': datetime.now().isoformat()
                }
                for i in range(len(chunks))
            ]

            # Add to collection
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=chunks,
                metadatas=metadatas
            )

            self.logger.info(f"Stored {len(chunks)} chunks for {metadata['file_name']}")
            return True

        except Exception as e:
            self.logger.error(f"Failed to store chunks: {e}")
            return False

    def process_all_documents(self, progress_callback: Optional[Callable[[int, int, str], None]] = None) -> Dict[str, any]:
        """
        Process all documents in the documents directory.

        Args:
            progress_callback: Optional callback function(current, total, message)

        Returns:
            Statistics dictionary:
            {
                'total_files': int,
                'processed_files': int,
                'failed_files': int,
                'total_chunks': int,
                'errors': List[str]
            }
        """
        stats = {
            'total_files': 0,
            'processed_files': 0,
            'failed_files': 0,
            'total_chunks': 0,
            'errors': []
        }

        # Scan for documents
        documents = self.scan_documents()
        stats['total_files'] = len(documents)

        if not documents:
            self.logger.warning("No documents found to process")
            return stats

        # Process each document
        for idx, doc_info in enumerate(documents, 1):
            file_name = doc_info['name']
            file_path = doc_info['path']

            if progress_callback:
                progress_callback(idx, len(documents), f"Processing {file_name}...")

            try:
                # Extract text
                text = self.extract_text(file_path)
                if not text:
                    error_msg = f"No text extracted from {file_name}"
                    self.logger.warning(error_msg)
                    stats['errors'].append(error_msg)
                    stats['failed_files'] += 1
                    continue

                # Chunk text
                chunks = self.chunk_text(text)
                if not chunks:
                    error_msg = f"No chunks created from {file_name}"
                    self.logger.warning(error_msg)
                    stats['errors'].append(error_msg)
                    stats['failed_files'] += 1
                    continue

                # Create embeddings
                embeddings = self.create_embeddings(chunks)
                if not embeddings:
                    error_msg = f"Failed to create embeddings for {file_name}"
                    self.logger.error(error_msg)
                    stats['errors'].append(error_msg)
                    stats['failed_files'] += 1
                    continue

                # Store in database
                metadata = {
                    'file_name': file_name,
                    'file_type': doc_info['extension']
                }

                if self.store_chunks(chunks, embeddings, metadata):
                    stats['processed_files'] += 1
                    stats['total_chunks'] += len(chunks)
                    self.logger.info(f"✅ Successfully processed {file_name} ({len(chunks)} chunks)")
                else:
                    error_msg = f"Failed to store chunks for {file_name}"
                    self.logger.error(error_msg)
                    stats['errors'].append(error_msg)
                    stats['failed_files'] += 1

            except Exception as e:
                error_msg = f"Error processing {file_name}: {str(e)}"
                self.logger.error(error_msg)
                stats['errors'].append(error_msg)
                stats['failed_files'] += 1

        # Final summary
        self.logger.info(f"Processing complete: {stats['processed_files']}/{stats['total_files']} files")
        return stats

    def get_stats(self) -> Dict[str, any]:
        """
        Get current database statistics.

        Returns:
            Dictionary with stats:
            {
                'total_chunks': int,
                'files_processed': int,
                'last_updated': str
            }
        """
        try:
            count = self.collection.count()

            # Get unique files
            if count > 0:
                results = self.collection.get()
                unique_files = set(meta['file_name'] for meta in results['metadatas'])

                # Get latest timestamp
                timestamps = [meta.get('processed_at', '') for meta in results['metadatas']]
                last_updated = max(timestamps) if timestamps else 'Never'
            else:
                unique_files = set()
                last_updated = 'Never'

            return {
                'total_chunks': count,
                'files_processed': len(unique_files),
                'last_updated': last_updated
            }

        except Exception as e:
            self.logger.error(f"Failed to get stats: {e}")
            return {
                'total_chunks': 0,
                'files_processed': 0,
                'last_updated': 'Error'
            }

    def clear_database(self) -> bool:
        """
        Clear all documents from the database.

        Returns:
            True if successful, False otherwise
        """
        try:
            # Delete the collection
            self.client.delete_collection(name=self.COLLECTION_NAME)

            # Recreate it
            self.collection = self.client.get_or_create_collection(
                name=self.COLLECTION_NAME,
                metadata={"description": "Interview preparation context documents"}
            )

            self.logger.info("Database cleared successfully")
            return True

        except Exception as e:
            self.logger.error(f"Failed to clear database: {e}")
            return False

    def query_similar(self, query_text: str, n_results: int = 5) -> List[Dict[str, any]]:
        """
        Query for similar documents.

        Args:
            query_text: Query text
            n_results: Number of results to return

        Returns:
            List of similar documents with metadata
        """
        if not OLLAMA_AVAILABLE:
            self.logger.error("Ollama not available for querying")
            return []

        try:
            # Create embedding for query
            response = ollama.embeddings(
                model=self.EMBEDDING_MODEL,
                prompt=query_text
            )
            query_embedding = response['embedding']

            # Query collection
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results
            )

            # Format results
            formatted_results = []
            for i in range(len(results['documents'][0])):
                formatted_results.append({
                    'document': results['documents'][0][i],
                    'metadata': results['metadatas'][0][i],
                    'distance': results['distances'][0][i] if 'distances' in results else None
                })

            return formatted_results

        except Exception as e:
            self.logger.error(f"Query failed: {e}")
            return []


# ============================================================================
# TESTING
# ============================================================================

if __name__ == "__main__":
    """
    Test the DocumentProcessor with sample files.
    """
    print("=" * 60)
    print("Document Processor Test")
    print("=" * 60)

    # Initialize processor
    processor = DocumentProcessor(
        documents_dir="../documents",
        db_path="../data/chroma_db"
    )

    print("\n📊 Current Database Stats:")
    stats = processor.get_stats()
    print(f"  Total chunks: {stats['total_chunks']}")
    print(f"  Files processed: {stats['files_processed']}")
    print(f"  Last updated: {stats['last_updated']}")

    print("\n🔍 Scanning for documents...")
    documents = processor.scan_documents()

    if not documents:
        print("  ⚠️  No documents found in ../documents/")
        print("  💡 Add some PDF, DOCX, TXT, or MD files to test!")
    else:
        print(f"  Found {len(documents)} documents:")
        for doc in documents:
            print(f"    - {doc['name']} ({doc['extension']}, {doc['size']} bytes)")

        print("\n🚀 Processing documents...")
        print("-" * 60)

        # Process with progress callback
        def progress_callback(current, total, message):
            print(f"  [{current}/{total}] {message}")

        results = processor.process_all_documents(progress_callback=progress_callback)

        print("-" * 60)
        print("\n✅ Processing Complete!")
        print(f"  Total files: {results['total_files']}")
        print(f"  Processed: {results['processed_files']}")
        print(f"  Failed: {results['failed_files']}")
        print(f"  Total chunks: {results['total_chunks']}")

        if results['errors']:
            print(f"\n⚠️  Errors encountered:")
            for error in results['errors']:
                print(f"    - {error}")

        # Show updated stats
        print("\n📊 Updated Database Stats:")
        stats = processor.get_stats()
        print(f"  Total chunks: {stats['total_chunks']}")
        print(f"  Files processed: {stats['files_processed']}")
        print(f"  Last updated: {stats['last_updated']}")

        # Test query if we have data
        if stats['total_chunks'] > 0:
            print("\n🔎 Testing similarity query...")
            test_query = "What are the key skills for product management?"
            results = processor.query_similar(test_query, n_results=3)

            if results:
                print(f"  Query: '{test_query}'")
                print(f"  Found {len(results)} similar chunks:")
                for i, result in enumerate(results, 1):
                    print(f"\n  Result {i}:")
                    print(f"    File: {result['metadata']['file_name']}")
                    print(f"    Chunk: {result['metadata']['chunk_id'] + 1}/{result['metadata']['total_chunks']}")
                    print(f"    Preview: {result['document'][:150]}...")

    print("\n" + "=" * 60)
    print("Test complete!")
    print("=" * 60)
